import re
from typing import Dict, List, Optional
import pdfplumber
from docx import Document
import os
from difflib import SequenceMatcher

# Funções utilitárias
def clean_text(text: Optional[str]) -> str:
    if not text:
        return ""
    return re.sub(r'\s+', ' ', text.replace("\n", " ").replace("\r", " ").replace(",", ".")).strip()

def extract_numbers(text: str) -> List[float]:
    matches = re.findall(r"[-+]?\d+(?:[.,]\d+)?", text)
    return [float(m.replace(",", ".")) for m in matches]

def is_header_line(line: str) -> bool:
    line_lower = line.lower()
    invalid_keywords = [
        "os resultados do teste apenas para referência", "cartão do relatório de análise", "nome: exemplo", "sexo: feminino", "idade: 31", "figura: peso padrão", "período do teste", "resultados reais do teste", "conselho de peritos", "real", "item de teste intervalo normal valor de medição real"
    ]
    return any(kw in line_lower for kw in invalid_keywords) or len(line) < 20 or not re.search(r'[a-zA-Z]{5,}', line)

def is_valid_name(name: str) -> bool:
    name_lower = name.lower()
    invalid_starts = ['(', ')', '-', 'do', 'da', 'de', 'e', 'o', 'a', 'ncia', 'gordo', 'grande', 'nível de']
    invalid_ends = [' de', ' do', ' da', ' ncia', ' o', '-', 'função', 'sistema', 'cartão do relatório']
    if name_lower.startswith(tuple(invalid_starts)) or name_lower.endswith(tuple(invalid_ends)):
        return False
    char_count = len(re.sub(r'[^a-zA-Z]', '', name))
    word_count = len(name.split())
    if char_count < 10 or char_count > 80 or word_count > 10:
        return False
    # Lista expandida de termos válidos do PDF
    valid_terms = [
        "viscosidade do sangue", "cristal de colesterol", "gordura do sangue", "resistência vascular", "elasticidade vascular", "demanda de sangue do miocárdio", "volume de perfusão do sangue do miocárdio", "consumo de oxigênio do miocárdio", "volume sistólico", "impedância do bombeamento de sangue do ventrículo esquerdo", "força de bombeamento efetiva do ventrículo esquerdo", "elasticidade da artéria coronária", "pressão de perfusão da artéria coronária", "elasticidade dos vasos sanguíneos do cérebro", "situação do fornecimento de sangue ao tecido cerebral", "coeficiente de secreção de pepsina", "coeficiente das funções peristálticas gástricas", "coeficiente das funções de absorção gástricas", "coeficiente das funções peristálticas do intestino delgado", "coeficiente das funções de absorção do intestino delgado", "metabolismo de proteínas", "função de produção de energia", "função de desintoxicação", "função de secreção de bílis", "teor de gordura do fígado", "globulina do soro sanguíneo (a/g)", "bilirrubina total (tbil)", "fosfatase alcalina (alp)", "ácido biliar total do soro sanguíneo (tba)", "bilirrubina (dbil)", "insulina", "polipeptídeo pancreático (pp)", "glucagon", "urobilinogênio", "ácido úrico", "nitrogênio uréico", "proteína urinária", "atividade pulmonar vc", "capacidade pulmonar total tlc", "resistência das vias aéreas ram", "teor de oxigênio no sangue arterial paco2", "fornecimento de sangue ao cérebro", "arterioesclerose cerebral", "condição das funções neurológicas", "indicador de depressão", "indicador de memória (zs)", "dimensão da protusão da fibra lombar", "grau de adesão da musculatura dos ombros", "limite de circulação dos membros", "grau de envelhecimento dos ligamentos", "coeficiente de oesteoclastos", "perda de cálcio", "grau de hiperplasia óssea", "grau de osteoporose", "densidade óssea", "calcificação coluna cervical", "calcificação coluna lombar", "coeficiente de hiperplasia óssea", "coeficiente de osteoporose", "coeficiente de reumatismo", "coeficiente de secreção de insulina", "coeficiente de açúcar no sangue", "coeficiente de açúcar na urina", "capacidade de reação fisica", "capacidade cerebral", "falta de água", "hipóxia", "ph", "bebida estimulante", "radiação eletromagnética", "tabaco/nicotina e outros", "resíduos tóxicos de pesticida", "cálcio", "ferro", "zinco", "selênio", "fósforo", "potássio", "magnésio", "cobre", "cobalto", "manganês", "iodo", "níquel", "flúor", "molibdênio", "vanádio", "estanho", "silício", "estrôncio", "boro", "estrogênio", "gonadotrofina", "prolactina", "progesterona", "coeficiente de vaginite", "coeficiente de inflamação pélvica", "coeficiente de anexite", "coeficiente de cervicite", "coeficiente de cisto nos ovarios", "índice dos radicais livres da pele", "índice de colágeno da pele", "índice de oleosidade da pele", "índice de imunidade da pele", "índice de hidratação da pele", "perda de hidratação da pele", "índice de dilatação dos vasos sanguíneos da pele", "índice de elasticidade da pele", "índice de melanina da pele", "índice de queratinócitos da pele", "índice de secreção da tireóide", "índice de secreção da paratireóide", "índice de secreção da glândula supra-renal", "índice de secreção da pituitária", "índice de secreção da glândula pineal", "índice de secreção do timo", "índice gonadal", "índice de linfonodo", "índice de imunidade das amígdalas", "índice da medula óssea", "índice do baço", "índice do timo", "índice de imunoglobulina", "índice de imunidade do trato respiratório", "índice de imunidade gastrointestinal", "índice de imunidade da mucosa", "coeficiente de fibrosidade da glândula mamária", "coeficiente de mastite aguda", "coeficiente de mastite crônica", "coeficiente de distúrbios endócrinos", "coeficiente de fibroadenoma", "vitamina a", "vitamina b1", "vitamina b2", "vitamina b3", "vitamina b6", "vitamina b12", "vitamina c", "vitamina d3", "vitamina e", "vitamina k", "lisina", "triptofano", "fenilalanina", "metionina", "treonina", "isoleucina", "leucina", "valina", "histidina", "arginina", "fosfatase alcalina óssea", "osteocalcina", "cartilagem grandes articulaçoes", "cartilagem pequenas articulações", "linha epifisária", "bolsas sob os olhos", "colágeno das rugas nos olhos", "pigmentação da pele (índice de olheiras)", "obstrução linfática", "afrouxamento e queda", "edema", "atividade das células dos olhos", "fadiga visual", "chumbo", "mercúrio", "cádmio", "crômio", "arsênico", "antimônio", "tálio", "alumínio", "índice de alergia a medicamentos", "índice de alergia álcool", "índice de alergia ao pólen", "índice de alergia antibioticos", "fibra química", "alergia a tintas e vernizes", "índice de alergia a poeira", "índice de alergia a fumos", "alergia a corante de tintas cabelo", "índice alergia de contato", "alergia a acessorios de metal", "índice alergia marisco", "índice alergia proteína do leite", "nicotinamida", "biotina", "ácido pantotênico", "ácido fólico", "coenzima q10", "glutationa", "coeficiente de metabolismo anormal de lipidos", "anormalidades tecido adiposo", "coeficiente de hiperinsulinemia", "coeficiente de anomalia hipotálamo núcleo", "coeficiente de conteúdo anormal de triglicerídeos", "olhos", "dentes", "cabelo e pele", "sistema endocrino", "circulação de sangue do coração e do cérebro", "estômago e intervalo intestinal", "sistema imunologico", "articulações", "tecido muscular", "metabolismo da gordura", "desintoxicação e metabolismo", "sistema reprodutivo", "sistema nervoso", "esqueleto", "coeficiente da função peristáltica do intestino grosso", "coeficiente de absorção do cólon", "coeficiente das bactérias intestinais (flora intestinal)", "coeficiente de pressão intraluminal", "tiroxina livre (t4)", "tiroglobulina", "os anticorpos antitireoglobulina", "triiodotironina (t3)", "ácido linoleico", "α-ácido linolênico", "γ-ácido linolênico", "ácido araquidônico", "estrogénio", "andrógeno", "progesterona(p)", "hormona luteinizante (lh)", "prolactina(prl)", "hormona estimuladora folícula (fsh)", "meridiano do pulmão tai yin da mão", "meridiano do intestino grosso yangming da mão", "meridiano do estômago yangming do pé", "meridiano baço/pancreas tai yn do pe", "meridiano do coração shao yin da mão", "meridiano do intestino delgado tai yang da mão", "meridiano da bexiga tai yang do pé", "meridiano dos rins shao yin do pé", "pericárdio", "triplo aquecedor shao yang da mão", "meridiano da vesícula biliar shao yang do pé", "meridiano do fígado jue yin do pé", "ren mai", "meridiano governador", "meridiano vital", "da mai", "índice de acidente vascular cerebral", "pulso (sv)", "resistência periférica do coração (trr)", "coeficiente da onda de pulso k", "saturação do oxigênio do sangue cerebrovascular (sa)", "volume do oxigênio do sangue cerebrovascular (caco2)", "pressão do oxigênio do sangue cerebrovascular (pao2)", "viscosidade do sangue", "colesterol total (tc)", "triglicerídeos (tg)", "lipoproteína de alta densidade (hdl-c)", "lipoproteína de baixa densidade (ldl-c)", "gordura neutra (mb)", "complexo imunológico circulatório (cic)", "hormona beta (folículo estimulante) fsh", "proteína de resposta", "fibrinogênio", "taxa de sedimentação", "índice imunitário de barreira tecidual", "índice de células imunitárias inatas", "índice de molécula imunitária inata", "índice imunitário celular", "índice de imunidade humoral", "vergonha", "culpa", "apatia", "dor", "medo", "desejo", "raiva", "orgulho", "coragem", "neutralidade", "vontade", "aceitação", "razão", "amor", "alegria", "paz", "iluminismo", "volume maré(vt)", "volume inspiratório (ti)", "capacidade residual funcional(frc)", "volume residual(rv)", "índice fosfolipídico", "índice esfingolípide", "índice de esfingomielilina", "índice de lecitina", "índice fosfolípide cerebral", "índice lipossômico", "índice de ácidos gordos saturados", "índice de ácidos gordos não saturados", "índice de ácidos gordos essenciais", "índice de triglicéridos"
    ]
    return any(term in name_lower for term in valid_terms)

def names_are_similar(a: str, b: str, threshold: float = 0.9) -> bool:
    return SequenceMatcher(None, normalize_name(a), normalize_name(b)).ratio() > threshold

def normalize_name(name: str) -> str:
    return re.sub(r'\s+', ' ', name.lower().replace("(", "").replace(")", "").strip())

# Função principal de extração (híbrido: tabelas + texto)
def extract_parameters_from_pdf(pdf_path: str) -> Dict[str, Dict[str, float]]:
    parameters = {}
    seen = set()
    buffer = ""

    table_settings = {
        "vertical_strategy": "lines",
        "horizontal_strategy": "text",
        "intersection_y_tolerance": 5,
    }

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            # Tenta extrair como tabela
            tables = page.extract_tables(table_settings)
            for table in tables:
                for row in table:
                    if len(row) < 3:
                        continue
                    name = clean_text(row[0])
                    range_str = clean_text(row[1])
                    val_str = clean_text(row[2])
                    if is_header_line(name + range_str):
                        continue
                    nums = extract_numbers(range_str + " " + val_str)
                    if len(nums) >= 3:
                        min_val, max_val, val = nums[:3]
                        if min_val < max_val and is_valid_name(name):
                            norm_name = normalize_name(name)
                            if any(names_are_similar(norm_name, s) for s in seen):
                                continue
                            seen.add(norm_name)
                            parameters[name] = {"min": min_val, "max": max_val, "valor": val}

            # Fallback para texto bruto se tabela falhar
            text = page.extract_text()
            if not text:
                continue
            lines = text.split("\n")
            for line in lines:
                line = clean_text(line)
                if not line or is_header_line(line):
                    buffer = ""
                    continue

                pattern = r"([A-Za-z\sKATEX_INLINE_OPEN/)]{10,80}?)\s*(\d+[.,]\d+)\s*-\s*(\d+[.,]\d+)\s*(\d+[.,]\d+)"
                match = re.search(pattern, line)
                if match:
                    raw_name, min_str, max_str, val_str = match.groups()
                    name = clean_text(buffer + " " + raw_name).strip()
                    if is_valid_name(name):
                        norm_name = normalize_name(name)
                        if any(names_are_similar(norm_name, s) for s in seen):
                            continue
                        seen.add(norm_name)
                        min_val = float(min_str.replace(",", "."))
                        max_val = float(max_str.replace(",", "."))
                        val = float(val_str.replace(",", "."))
                        if min_val < max_val:
                            parameters[name] = {"min": min_val, "max": max_val, "valor": val}
                    buffer = ""
                else:
                    if len(buffer + " " + line) < 50 and is_valid_name(line):
                        buffer += " " + line
                    else:
                        buffer = ""

    return parameters

# Validação com tolerância
def validate_parameters(parameters: Dict[str, Dict[str, float]]) -> List[Dict[str, any]]:
    anomalies = []
    for name, data in parameters.items():
        val = data.get("valor")
        min_val = data.get("min")
        max_val = data.get("max")
        if val is None or min_val is None or max_val is None or min_val >= max_val:
            continue
        # Tolerância de 0.001 para floats
        if not (min_val - 0.001 <= val <= max_val + 0.001):
            status = "Abaixo" if val < min_val else "Acima"
            anomalies.append({
                "item": name,
                "valor_real": val,
                "status": status,
                "normal_min": min_val,
                "normal_max": max_val,
            })
    return anomalies

# Geração de relatório (inalterada)
def generate_report(anomalies: List[Dict[str, any]], therapist: str, registry: str, output_path: str) -> None:
    doc = Document()
    doc.add_heading("Relatório de Anomalias", level=1)
    doc.add_paragraph(f"Terapeuta: {therapist}   Registro: {registry}")

    if not anomalies:
        doc.add_paragraph("🎉 Todos os parâmetros dentro da normalidade.")
    else:
        doc.add_paragraph(f"⚠️ {len(anomalies)} anomalias encontradas:")
        for a in anomalies:
            doc.add_paragraph(
                f"• {a['item']}: {a['valor_real']:.3f} "
                f"({a['status']} do normal; Normal: {a['normal_min']}–{a['normal_max']})"
            )
    doc.save(output_path)
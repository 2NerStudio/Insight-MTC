import pdfplumber
import re
from docx import Document
from typing import Dict, Tuple, List, Union

def extract_patient_info(text: str) -> Dict[str, str]:
    """Extracts patient information from the header"""
    patient_info = {
        'nome': 'N/A',
        'sexo': 'N/A',
        'idade': 'N/A',
        'data_exame': 'N/A'
    }
    
    # Extract name, gender and age
    name_match = re.search(r"Nome:\s*(.+?)\s*\|Sexo:\s*(.+?)\s*\|Idade:\s*(\d+)", text)
    if name_match:
        patient_info.update({
            'nome': name_match.group(1).strip(),
            'sexo': name_match.group(2).strip(),
            'idade': name_match.group(3).strip()
        })
    
    # Extract exam date
    date_match = re.search(r"Período do teste:\s*(\d{4}/\d{2}/\d{2}\s*\d{2}:\d{2})", text)
    if date_match:
        patient_info['data_exame'] = date_match.group(1).strip()
    
    return patient_info

def clean_text(text: str) -> str:
    """Cleans and normalizes text for processing"""
    text = re.sub(r'\n+', ' ', text)  # Replace multiple newlines with space
    text = re.sub(r'\s+', ' ', text)  # Replace multiple spaces with single space
    text = re.sub(r'-\s+', '', text)  # Join hyphenated words
    return text.strip()

def extract_exam_data(text: str) -> Dict[str, Dict[str, Union[Tuple[float, float], float, str]]]:
    """Extracts exam parameters and values from the text"""
    data = {'parametros': {}, 'valores': {}}
    
    # Improved pattern to match exam items
    pattern = re.compile(
        r'(?P<sistema>.+?)\s*\n'  # System (may be on previous line)
        r'(?P<item>.+?)\s+'       # Test item
        r'(?P<intervalo>\d+[\.,]\d+\s*[-–]\s*\d+[\.,]\d+)\s+'  # Reference range
        r'(?P<valor>\d+[\.,]\d+)\s*'  # Measured value
        r'(?P<conselho>.*?)(?=\n\s*[A-ZÀ-Ú]|\Z)',  # Advice (until next item or end)
        re.MULTILINE | re.DOTALL
    )
    
    for match in pattern.finditer(text):
        system = clean_text(match.group('sistema'))
        item = clean_text(match.group('item'))
        range_vals = match.group('intervalo').replace(',', '.').replace(' ', '')
        value = match.group('valor').replace(',', '.')
        advice = clean_text(match.group('conselho'))
        
        try:
            min_val, max_val = map(float, re.split(r'[-–]', range_vals))
            float_value = float(value)
            
            key = f"{system} | {item}"
            data['parametros'][key] = (min_val, max_val, advice)
            data['valores'][key] = float_value
        except (ValueError, TypeError):
            continue
    
    return data

def analyze_results(data: Dict) -> Dict:
    """Analyzes results and identifies anomalies"""
    anomalies = []
    normal = []
    
    for key, value in data['valores'].items():
        if key in data['parametros']:
            min_val, max_val, advice = data['parametros'][key]
            
            status = "DENTRO"
            if value < min_val:
                status = "ABAIXO"
            elif value > max_val:
                status = "ACIMA"
            
            result = {
                'parametro': key,
                'valor': value,
                'intervalo': f"{min_val:.3f} - {max_val:.3f}",
                'status': status,
                'conselho': advice if status != "DENTRO" else ""
            }
            
            if status != "DENTRO":
                anomalies.append(result)
            else:
                normal.append(result)
    
    return {
        'anomalias': anomalies,
        'normais': normal,
        'total_parametros': len(data['valores']),
        'total_anomalias': len(anomalies)
    }

def create_report(data: Dict, analysis: Dict, output_path: str = "Relatorio_Exames.docx") -> str:
    """Creates a comprehensive report in Word format"""
    doc = Document()
    
    # Header
    doc.add_heading('RELATÓRIO DE ANÁLISE DE EXAMES', level=1)
    
    # Patient information
    patient = data.get('paciente', {})
    doc.add_paragraph(f"Paciente: {patient.get('nome', 'N/A')}")
    doc.add_paragraph(f"Sexo: {patient.get('sexo', 'N/A')} | Idade: {patient.get('idade', 'N/A')}")
    doc.add_paragraph(f"Data do Exame: {patient.get('data_exame', 'N/A')}")
    doc.add_paragraph("\n")
    
    # Summary
    doc.add_heading('RESUMO', level=2)
    doc.add_paragraph(f"Total de Parâmetros Analisados: {analysis['total_parametros']}")
    doc.add_paragraph(f"Anomalias Detectadas: {analysis['total_anomalias']}")
    doc.add_paragraph(f"Percentual de Anomalias: {analysis['total_anomalias']/analysis['total_parametros']:.1%}")
    doc.add_paragraph("\n")
    
    # Anomalies table
    if analysis['anomalias']:
        doc.add_heading('PARÂMETROS COM ANOMALIAS', level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Header
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Parâmetro'
        hdr_cells[1].text = 'Valor'
        hdr_cells[2].text = 'Intervalo Normal'
        hdr_cells[3].text = 'Status'
        hdr_cells[4].text = 'Recomendações'
        
        # Add rows
        for item in analysis['anomalias']:
            row_cells = table.add_row().cells
            row_cells[0].text = item['parametro']
            row_cells[1].text = f"{item['valor']:.3f}"
            row_cells[2].text = item['intervalo']
            row_cells[3].text = item['status']
            row_cells[4].text = item['conselho']
    
    doc.add_paragraph("\n")
    
    # Normal results summary
    if analysis['normais']:
        doc.add_heading('PARÂMETROS NORMAIS', level=2)
        doc.add_paragraph(f"Total de parâmetros dentro da normalidade: {len(analysis['normais'])}")
    
    # Footer
    doc.add_paragraph("\n\n")
    doc.add_paragraph("Este relatório foi gerado automaticamente com base nos dados extraídos do exame.")
    doc.add_paragraph("Os resultados devem ser interpretados por um profissional de saúde qualificado.")
    
    doc.save(output_path)
    return output_path

def process_pdf_report(pdf_path: str) -> str:
    """Main function to process the PDF and generate report"""
    try:
        # Extract data from PDF
        with pdfplumber.open(pdf_path) as pdf:
            full_text = "\n".join([page.extract_text() for page in pdf.pages])
        
        patient_data = extract_patient_info(full_text)
        exam_data = extract_exam_data(full_text)
        
        if not exam_data['valores']:
            raise ValueError("Não foi possível extrair dados dos exames do PDF.")
        
        exam_data['paciente'] = patient_data
        analysis = analyze_results(exam_data)
        
        # Generate report filename based on patient name
        report_name = f"Relatorio_{patient_data['nome'].replace(' ', '_')}.docx"
        report_path = create_report(exam_data, analysis, report_name)
        
        print(f"Relatório gerado com sucesso: {report_path}")
        return report_path
    
    except Exception as e:
        print(f"Erro ao processar o relatório: {str(e)}")
        raise

if __name__ == "__main__":
    import sys
    if len(sys.argv) != 2:
        print("Uso: python app.py <caminho_do_pdf>")
        sys.exit(1)
    
    pdf_file = sys.argv[1]
    try:
        report_file = process_pdf_report(pdf_file)
        print(f"Relatório salvo como: {report_file}")
    except Exception as e:
        print(f"Falha ao gerar relatório: {e}")
        sys.exit(1)
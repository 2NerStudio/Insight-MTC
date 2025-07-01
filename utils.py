from docx import Document
from io import BytesIO
from PyPDF2 import PdfReader

def transformar_relatorio(arquivo, nome_terapeuta, registro):
    if arquivo.name.endswith(".pdf"):
        pdf = PdfReader(arquivo)
        texto = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
    elif arquivo.name.endswith(".txt"):
        texto = arquivo.read().decode("utf-8")
    elif arquivo.name.endswith(".docx"):
        import docx
        doc = docx.Document(arquivo)
        texto = "\n".join([p.text for p in doc.paragraphs])
    else:
        texto = ""

    # Placeholder: IA ainda não está ativa
    rodape = f"\n\n---\nRelatório elaborado por {nome_terapeuta} — Registro: {registro}\n"
    return "[Texto traduzido com linguagem MTC]\n\n" + texto + rodape

def exportar_para_docx(texto):
    doc = Document()
    doc.add_heading("Relatório Traduzido pela MTC", level=1)
    doc.add_paragraph(texto)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
